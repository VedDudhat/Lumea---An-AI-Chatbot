from tika import parser
from docx import Document
import re
import os
from embedding_model import Model
import warnings
warnings.filterwarnings("ignore")


class pipe:
    def __init__(self):
        self.model = Model()

    def pass_resume_text(self,text):
        """
        The function takes String as input aand store the string in vectore store

        Args:
            text (str): content of a single file that need to be store in vectore store
        
        Returns:
            bool: if file is succesfully added to the vector database
        """
        content = text.lower()
        content = re.sub(r'\s',' ',content).strip()
        return self.model.store_in_vectordatabase(input=content)

    def pass_resume(self,input_path):
        """
        Detect whether the input is a PDF, DOCX, or a folder containing PDFs/DOCXs.
        Raises an error if the input is an unsupported file type.
        
        Args:
            input_path (str): Path to the file or folder.
        
        Returns:
            bool: if file is succesfully added to the vector database
        """

        if not os.path.exists(input_path):
            raise FileNotFoundError("Error: The specified path does not exist.")
        
        if os.path.isfile(input_path):
            ext = os.path.splitext(input_path)[1].lower()
            if ext == ".pdf":
                print("Single PDF file detected.")
                parsed_pdf = parser.from_file(input_path)
                content = parsed_pdf['content']
                content = content.lower()
                content = re.sub(r'\s',' ',content).strip()
                return self.model.store_in_vectordatabase(input=content)
            
            elif ext == ".docx":
                print("Single DOCX file detected.")
                doc = Document(input_path)
                data = "".join([para.text for para in doc.paragraphs])
                data = data.lower()
                data = re.sub(r'\s+',' ',data).strip()   
                return self.model.store_in_vectordatabase(input=data)
            
            else:
                raise ValueError(f"Error: Unsupported file type '{ext}'. Only PDF and DOCX are allowed.")
        
        elif os.path.isdir(input_path):
            files = os.listdir(input_path)
            valid_files = [f for f in files if f.lower().endswith(('.pdf', '.docx'))]
            
            if not valid_files:
                raise ValueError("Error: No valid PDF or DOCX files found in the folder.")
            
            print("Multiple supported file detected in folder.")

            for f in valid_files:
                input_path_2 = input_path + "\\" + f
                if f.lower().endswith(".pdf"):
                    parsed_pdf = parser.from_file(input_path_2)
                    content = parsed_pdf['content']
                    content = content.lower()
                    content = re.sub(r'\s',' ',content).strip()
                    self.model.store_in_vectordatabase(input=content)
                
                elif f.lower().endswith(".docx"):
                    doc = Document(input_path_2)
                    data = "".join([para.text for para in doc.paragraphs])
                    data = data.lower()
                    data = re.sub(r'\s+',' ',data).strip()   
                    self.model.store_in_vectordatabase(input=data)
                
                else:
                    raise ValueError(f"Error: Unsupported file type '{ext}'. Only PDF and DOCX are allowed.")
            return True
        
        else:
            raise ValueError("Error: Invalid input type.")
        
    def pass_temp_resume(self,input):
        """
        Take the string as input and save it into temporary vectore database
        
        Args:
            input (dict): dictionary with filename as key and string as value that need to be added in temporary vector database
        
        Returns:
            bool: if file is succesfully added to the temporary vector database
        """
        flag = []
        self.model.delete_temp_database()
        for key,value in input.items():
            content = value.lower()
            content = re.sub(r'\s',' ',content).strip()
            flag.append(self.model.store_in_temp_vectordatabase(input=content))
        if False in flag:
            return False
        else:
            return True
