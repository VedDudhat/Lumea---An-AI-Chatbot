import os
import streamlit as st
from tika import parser
from docx import Document
import tempfile
import time
import zipfile
from self_data_recruiter_llm import Agent
from langgraph.checkpoint.memory import InMemorySaver
from data_pipeline import pipe
import warnings
warnings.filterwarnings("ignore")

st.set_page_config(
    page_title="Self Data Recruiter",
    initial_sidebar_state="expanded",
)
hide_st_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    </style>
"""
st.markdown(hide_st_style, unsafe_allow_html=True)
if "current_page" not in st.session_state:
    st.session_state.current_page = "self_data_recruiter"

if st.session_state.current_page != "self_data_recruiter":
    st.session_state.clear()
    st.session_state.current_page = "self_data_recruiter"

if "files" not in st.session_state:
    st.session_state.files = []

if "folder" not in st.session_state:
    st.session_state.folder = []

def response_generator(result):
    for line in result.split("\n"):
        words = line.split()
        for word in words:
            yield word + " "
            time.sleep(0.05)
        yield "\n"

def extract_text(uploaded_file, file_type):
    if file_type == "pdf":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(uploaded_file.getvalue())
            temp_file_path = temp_file.name
        
        doc = parser.from_file(temp_file_path)
        os.remove(temp_file_path)
        return { uploaded_file.name.rsplit(".",1)[0] : doc.get("content", "No text found") }
    
    elif file_type == "docx":
        doc = Document(uploaded_file)
        return { uploaded_file.name.rsplit(".",1)[0] : "".join([para.text for para in doc.paragraphs])} 

    elif file_type == "zip":
        return extract_text_from_zip(uploaded_file)

    else:
        return "Unsupported file type"
    
def extract_text_from_zip(uploaded_zip):
    extracted_texts = {}
    
    with tempfile.TemporaryDirectory() as temp_dir:
        zip_path = os.path.join(temp_dir, "uploaded.zip")
        with open(zip_path, "wb") as f:
            f.write(uploaded_zip.getvalue())
        
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
            
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    
                    if file.endswith(".pdf"):
                        doc = parser.from_file(file_path)
                        extracted_texts[file.rsplit(".",1)[0]] = doc.get("content", "No text found")
                    elif file.endswith(".docx"):
                        doc = Document(file_path)
                        extracted_texts[file.rsplit(".",1)[0]] = "".join([para.text for para in doc.paragraphs])
    
    return extracted_texts

    
st.markdown("""
<style>
html, body, [data-testid="stAppViewContainer"], [data-testid="stApp"] {
    height: auto;
    overflow: hidden;
    background: #070707;
}
body {
    background: transparent;
    margin: 0;
    padding: 0;
    overflow: hidden;
}
@keyframes slideUp {
  from { transform: translateY(0); opacity: 1; }
  to { transform: translateY(-100%); opacity: 0; height: 0; }
}
@keyframes gradient {
    0% {
        background-position: 0% 0%;
    }
    50% {
        background-position: 100% 100%;
    }
    100% {
        background-position: 0% 0%;
    }
}
[data-testid="stApp"] {
    z-index: auto;
    width: auto;
    height: auto;
    display: flex;
    top: 0px;
    justify-content: center;
    align-items: center;
    color: white;
    opacity: 1;
}

:root {
  --card-height: 100%;
  --card-width: 100%;
  --color1: 0, 0, 252;
  --color2: 74, 92, 255;
  --color3: 0, 157, 255;
  --color4: 0 ,0, 255;
  --color5:  74, 92, 255;
  --color-interactive: 140, 100, 255;
  --blending: hard-light;
}
@keyframes moveInCircle {
  0% {
    transform: rotate(0deg);
  }
  50% {
    transform: rotate(180deg);
  }
  100% {
    transform: rotate(360deg);
  }
}
            
@keyframes moveAlongTopEdge {
  0% {
    transform: translateX(0%) translateY(0%);
  }
  50% {
    transform: translateX(100%) translateY(0%);
  }
  100% {
    transform: translateX(0%) translateY(0%);
  }
}
            
@keyframes moveAlongBottomEdge {
  0% {
    transform: translateX(100%) translateY(0%);
  }
  50% {
    transform: translateX(0%) translateY(0%);
  }
  100% {
    transform: translateX(100%) translateY(0%);
  }
}
            
@keyframes moveAlongLeftEdge {
  0% {
    transform: translateX(0%) translateY(0%);
  }
  50% {
    transform: translateX(0%) translateY(100%);
  }
  100% {
    transform: translateX(0%) translateY(0%);
  }
}
            
@keyframes moveAlongRightEdge {
  0% {
    transform: translateX(0%) translateY(100%);
  }
  50% {
    transform: translateX(0%) translateY(0%);
  }
  100% {
    transform: translateX(0%) translateY(100%);
  }
}
            
.gradient-bg {
    position: fixed;
    top: 0;
    left: 0;
    width: 100%;
    height: 100%;
    z-index: 0;
    overflow: hidden;
}
            
svg {
    position: relative;
    width: 100%;
    height: 100%;
    top:0;
    left:0;
}
            
.gradients-container {
    filter: url(#goo) blur(40px) ;
    width: 100%;
    height: 100%;
    z-index:1;
}

.g1 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color1), 0.8) 0, rgba(var(--color1), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 50%;
    height: 25%;
    top: 0;
    left: 0;

    transform-origin: top left;
    animation: moveAlongTopEdge 20s linear infinite;
    opacity: 1;
}
            
.g2 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color2), 0.8) 0, rgba(var(--color2), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 50%;
    height: 30%;
    top: 0;
    right: 0;

    transform-origin: top right;
    animation: moveInCircle 20s reverse infinite;

    opacity: 1;
}

.g3 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color3), 0.8) 0, rgba(var(--color3), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 40%;
    height: 20%;
    bottom: 0;
    left: 0;

    transform-origin: bottom left;
    animation: moveInCircle 20s linear infinite;

    opacity: 1;
}

.g4 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color4), 0.8) 0, rgba(var(--color4), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 45%;
    height: 40%;
    bottom: 0;
    right: 0;

    transform-origin: bottom right;
    animation: moveHorizontal 20s reverse infinite;

    opacity: 0.7;
}

.g6 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color4), 0.8) 0, rgba(var(--color4), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 60%;
    height: 40%;
    top: 0;
    right: 0;

    transform-origin: top right;
    animation: moveAlongRightEdge 20s linear infinite;

    opacity: 1;
}
.g7 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color2), 0.8) 0, rgba(var(--color2), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 50%;
    height: 40%;
    bottom : 0;
    left: 0;

    transform-origin: bottom left;
    animation: moveAlongLeftEdge 20s reverse infinite;

    opacity: 1;
}
.g8 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color3), 0.8) 0, rgba(var(--color3), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 50%;
    height: 40%;
    bottom : 0;
    right: 0;

    transform-origin: bottom right;
    animation: moveAlongBottomEdge 20s linear infinite;

    opacity: 1;
}
            
.glass-box {
    position: relative;
    border-radius: 10px;
    padding: 2rem;
    backdrop-filter: blur(20px);
    -webkit-backdrop-filter: blur(100px);
    border: 1px solid rgba(255, 255, 255, 0.1);
    box-shadow: 0 0 10px #FF9131, 0 0 40px #6736F3, 0 0 80px #2196f3;
    justify-content: center;
    align-items: center;
    text-align: center;
    z-index:0;
}
                          
[class="st-emotion-cache-hzygls eht7o1d3"]{
    background-color:transparent;       
}
            
.header {
  transition: all 2s ease-in-out;
  text-align: center;
  padding-top: 100px;
}
            
.hidden-h1 {
  animation: slideUp 3s forwards;
  overflow: hidden;
  height: 0;
  opacity: 0;
}
   
</style>
            
<div class="gradient-bg">
    <svg xmlns="http://www.w3.org/2000/svg>
    <defs>
      <filter id="goo">
        <feGaussianBlur in="SourceGraphic" stdDeviation="10" result="blur" />
        <feColorMatrix in="blur" mode="matrix" values="1 0 0 0 0  0 1 0 0 0  0 0 1 0 0  0 0 0 18 -8" result="goo" />
        <feBlend in="SourceGraphic" in2="goo" />
      </filter>
    </defs>
    </svg>
  <div class="gradients-container">
    <div class="g1"></div>
    <div class="g2"></div>
    <div class="g3"></div>
    <div class="g4"></div>
    <div class="g5"></div>
    <div class="g6"></div>
    <div class="g7"></div>
    <div class="g8"></div>
  </div>                     
</div>       
""", unsafe_allow_html=True)
st.markdown(
    """
    <style>
    [data-testid="stSidebar"] > div:first-child {
        
      height:100%;
      width: 100%;
      background-color: #000000; !important;
      backdrop-filter: blur(20px) !important;
    }   
    </style>
    """,
    unsafe_allow_html=True)
st.markdown("""
<div class="header">
    <div class="hidden-header">
        <div class="glass-box"><h1>Welcome to Self Data Recruiter</h1> </br>
            <h5> This is the recruiter side of our App where you can use your resume database and then have talk with Lumea"</h5>
        </div>             
    </div>
</div></br>""", unsafe_allow_html=True)

uploaded_file = st.sidebar.file_uploader("Upload a resume or zip folder containing resumes", type=["docx","pdf","zip"])

if uploaded_file is None:
    st.warning("⚠️ Please upload a file to proceed.")
    st.stop()


if uploaded_file is not None:
    if uploaded_file.name not in st.session_state.folder:
        st.session_state.folder.append(uploaded_file.name)
        with st.spinner("Processing the file...", show_time=True):
            file_type = uploaded_file.name.split(".")[-1]
            text = extract_text(uploaded_file, file_type)
        with st.spinner("Creating temporary vector store", show_time=True):
            Pipe = pipe()
            flag = Pipe.pass_temp_resume(text)
            if flag:
                for key,value in text.items():
                    st.session_state.files.append(key)
                st.sidebar.success(f"{len(st.session_state.files)} document successfully uploaded to vector store")
            else:
                st.sidebar.warning("error in uploading files to vector store")
        if 'messages' not in st.session_state:
            st.session_state.messages = []
            st.session_state.memory = InMemorySaver()
            st.session_state.thread = {"configurable": {"thread_id": "1"}}
            st.session_state.abot = Agent(st.session_state.memory)
            print("Agent variable initialized")

abot = st.session_state.abot

for messages in st.session_state.messages:
    with st.chat_message(messages["role"]):
        st.markdown(messages["content"])

if prompt:= st.chat_input("Ask your Question"):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.spinner("Generating response", show_time=True):
        result = abot.graph.invoke({"user_input":[prompt]}, st.session_state.thread)
    
    with st.chat_message("assistant"):
        response = st.write_stream(response_generator(result['msg'][-1].content))

    st.session_state.messages.append({"role": "assistant", "content": response})
