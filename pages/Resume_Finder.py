import streamlit as st 
from recruiter_llm import Agent
from langgraph.checkpoint.memory import InMemorySaver
import time
from tika import parser
import tempfile
import zipfile
from docx import Document
import os
from data_pipeline import pipe
import warnings
warnings.filterwarnings("ignore")


st.set_page_config(
    page_title="Recruiter Page",
    initial_sidebar_state="expanded",
)
hide_st_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    </style>
"""
st.markdown(hide_st_style, unsafe_allow_html=True)
def response_generator(result):
    for line in result.split("\n"):
        words = line.split()
        for word in words:
            yield word + " "
            time.sleep(0.05)
        yield "\n"

def extract_text(uploaded_file, file_type):
    if file_type == "pdf":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(uploaded_file.getvalue())
            temp_file_path = temp_file.name
        
        doc = parser.from_file(temp_file_path)
        os.remove(temp_file_path)
        return { uploaded_file.name.rsplit(".",1)[0] : doc.get("content", "No text found") }
    
    elif file_type == "docx":
        doc = Document(uploaded_file)
        return { uploaded_file.name.rsplit(".",1)[0] : "".join([para.text for para in doc.paragraphs])} 

    elif file_type == "zip":
        return extract_text_from_zip(uploaded_file)

    else:
        return "Unsupported file type"
    
def extract_text_from_zip(uploaded_zip):
    extracted_texts = {}
    
    with tempfile.TemporaryDirectory() as temp_dir:
        zip_path = os.path.join(temp_dir, "uploaded.zip")
        with open(zip_path, "wb") as f:
            f.write(uploaded_zip.getvalue())
        
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
            
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    
                    if file.endswith(".pdf"):
                        doc = parser.from_file(file_path)
                        extracted_texts[file.rsplit(".",1)[0]] = doc.get("content", "No text found")
                    elif file.endswith(".docx"):
                        doc = Document(file_path)
                        extracted_texts[file.rsplit(".",1)[0]] = "".join([para.text for para in doc.paragraphs])
    
    return extracted_texts

if 'current_page' not in st.session_state:
    st.session_state.current_page = "recruiter_page"

if st.session_state.current_page != "recruiter_page":
    st.session_state.clear()

# st.sidebar.header("This is the recruiter side of our App and here you can retrieve the resumes and ask question about the resume to our chatbot Stella")
st.markdown("""
<style>
html, body, [data-testid="stAppViewContainer"], [data-testid="stApp"] {
    height: auto;
    overflow: hidden;
    background: #070707;
}
body {
    margin: 0;
    padding: 0;
    overflow: hidden;
}
@keyframes gradient {
    0% {
        background-position: 0% 0%;
    }
    50% {
        background-position: 100% 100%;
    }
    100% {
        background-position: 0% 0%;
    }
}
@keyframes slideUp {
  0% {
    transform: translateY(0%);
    opacity: 1;
  }
  100% {
    transform: translateY(-100%);
    opacity: 0;
    height: 0;
    visibility: hidden;
  }
}
.slide-up {
  animation: slideUp 1s forwards;
}
[data-testid="stApp"] {
    z-index: auto;
    width: auto;
    height: auto;
    display: flex;
    top: 0px;
    justify-content: center;
    align-items: center;
    color: white;
    opacity: 1;
}

:root {
  --card-height: 100%;
  --card-width: 100%;
  --color1: 0, 0, 252;
  --color2: 74, 92, 255;
  --color3: 0, 157, 255;
  --color4: 0 ,0, 255;
  --color5:  74, 92, 255;
  --color-interactive: 140, 100, 255;
  --blending: hard-light;
}
@keyframes moveInCircle {
  0% {
    transform: rotate(0deg);
  }
  50% {
    transform: rotate(180deg);
  }
  100% {
    transform: rotate(360deg);
  }
}
            
@keyframes moveAlongTopEdge {
  0% {
    transform: translateX(0%) translateY(0%);
  }
  50% {
    transform: translateX(100%) translateY(0%);
  }
  100% {
    transform: translateX(0%) translateY(0%);
  }
}
            
@keyframes moveAlongBottomEdge {
  0% {
    transform: translateX(100%) translateY(0%);
  }
  50% {
    transform: translateX(0%) translateY(0%);
  }
  100% {
    transform: translateX(100%) translateY(0%);
  }
}
            
@keyframes moveAlongLeftEdge {
  0% {
    transform: translateX(0%) translateY(0%);
  }
  50% {
    transform: translateX(0%) translateY(100%);
  }
  100% {
    transform: translateX(0%) translateY(0%);
  }
}
            
@keyframes moveAlongRightEdge {
  0% {
    transform: translateX(0%) translateY(100%);
  }
  50% {
    transform: translateX(0%) translateY(0%);
  }
  100% {
    transform: translateX(0%) translateY(100%);
  }
}
            
.gradient-bg {
    position: fixed;
    top: 0;
    left: 0;
    width: 100%;
    height: 100%;
    z-index: 0;
    overflow: hidden;
}
            
svg {
    position: relative;
    width: 100%;
    height: 100%;
    top:0;
    left:0;
}
            
.gradients-container {
    filter: url(#goo) blur(40px) ;
    width: 100%;
    height: 100%;
    z-index:1;
}

.g1 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color1), 0.8) 0, rgba(var(--color1), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 50%;
    height: 25%;
    top: 0;
    left: 0;

    transform-origin: top left;
    animation: moveAlongTopEdge 20s linear infinite;
    opacity: 1;
}
            
.g2 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color2), 0.8) 0, rgba(var(--color2), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 50%;
    height: 30%;
    top: 0;
    right: 0;

    transform-origin: top right;
    animation: moveInCircle 20s reverse infinite;

    opacity: 1;
}

.g3 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color3), 0.8) 0, rgba(var(--color3), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 40%;
    height: 20%;
    bottom: 0;
    left: 0;

    transform-origin: bottom left;
    animation: moveInCircle 20s linear infinite;

    opacity: 1;
}

.g4 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color4), 0.8) 0, rgba(var(--color4), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 45%;
    height: 40%;
    bottom: 0;
    right: 0;

    transform-origin: bottom right;
    animation: moveHorizontal 20s reverse infinite;

    opacity: 0.7;
}

.g6 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color4), 0.8) 0, rgba(var(--color4), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 60%;
    height: 40%;
    top: 0;
    right: 0;

    transform-origin: top right;
    animation: moveAlongRightEdge 20s linear infinite;

    opacity: 1;
}
.g7 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color2), 0.8) 0, rgba(var(--color2), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 50%;
    height: 40%;
    bottom : 0;
    left: 0;

    transform-origin: bottom left;
    animation: moveAlongLeftEdge 20s reverse infinite;

    opacity: 1;
}
.g8 {
    position: absolute;
    background: radial-gradient(circle at center, rgba(var(--color3), 0.8) 0, rgba(var(--color3), 0) 50%) no-repeat;
    mix-blend-mode: var(--blending);

    width: 50%;
    height: 40%;
    bottom : 0;
    right: 0;

    transform-origin: bottom right;
    animation: moveAlongBottomEdge 20s linear infinite;

    opacity: 1;
}
            
.glass-box {
    position: relative;
    border-radius: 10px;
    padding: 2rem;
    backdrop-filter: blur(20px);
    -webkit-backdrop-filter: blur(100px);
    border: 1px solid rgba(255, 255, 255, 0.1);
    box-shadow: 0 0 10px #FF9131, 0 0 40px #6736F3, 0 0 80px #2196f3;
    justify-content: center;
    align-items: center;
    text-align: center;
    z-index:0;
}
                          
[class="st-emotion-cache-hzygls eht7o1d3"]{
    background-color:transparent;       
}
[data-testid="stChatInput"]{
  width: 100%;
  background: #000000;
  border-radius: 20px;
  box-shadow: 0 0 30px #00bfff;
  overflow: hidden;         
}            
.header {
  transition: all 2s ease-in-out;
  text-align: center;
  padding-top: 100px;
}
            
.hidden-h1 {
  animation: slideUp 3s forwards;
  overflow: hidden;
  height: 0;
  opacity: 0;
}
   
</style>
           
<div class="gradient-bg">
    <svg xmlns="http://www.w3.org/2000/svg>
    <defs>
      <filter id="goo">
        <feGaussianBlur in="SourceGraphic" stdDeviation="10" result="blur" />
        <feColorMatrix in="blur" mode="matrix" values="1 0 0 0 0  0 1 0 0 0  0 0 1 0 0  0 0 0 18 -8" result="goo" />
        <feBlend in="SourceGraphic" in2="goo" />
      </filter>
    </defs>
    </svg>
  <div class="gradients-container">
    <div class="g1"></div>
    <div class="g2"></div>
    <div class="g3"></div>
    <div class="g4"></div>
    <div class="g5"></div>
    <div class="g6"></div>
    <div class="g7"></div>
    <div class="g8"></div>
  </div>                     
</div>       
""", unsafe_allow_html=True)
st.markdown(
    """
    <style>
    [data-testid="stSidebar"] > div:first-child {
        
      height:100%;
      width: 100%;
      background-color: #000000; !important;
      backdrop-filter: blur(20px) !important;
    }   
    </style>
    """,
    unsafe_allow_html=True)
st.markdown("""
<div class="header">
    <div class="hidden-header">
        <div class="glass-box"><h1>Welcome to Resume Finder</h1> </br>
            <h5> This is the recruiter side of our App and here you can retrieve the resumes and ask question about the resume to our chatbot Lumea"</h5>
        </div>             
    </div>
</div></br>""", unsafe_allow_html=True)

st.sidebar.write("if you want to add your resume to the vectore store, please upload it here")
uploaded_file = st.sidebar.file_uploader("Upload a file", type=["pdf", "docx", "zip"])

if "files" not in st.session_state:
    st.session_state.files = []

if uploaded_file is not None:
    if uploaded_file.name not in st.session_state.files:
        st.session_state.files.append(uploaded_file.name)
        with st.spinner("Uploading the file",show_time=True):
            Pipe = pipe()
            file_type = uploaded_file.name.split(".")[-1]
            if file_type == "pdf":
                st.sidebar.write("PDF file uploaded")
                text = extract_text(uploaded_file,file_type)
                for key,value in text.items():
                    flag = Pipe.pass_resume_text(value)
                if flag:
                    st.sidebar.success("document successfully uploaded to vector store")
                else:
                    st.sidebar.warning("error in uploadinf file to vector store")

            elif file_type == "docx":
                st.sidebar.write("DOCX file uploaded")
                text = extract_text(uploaded_file,file_type)
                for key,value in text.items():
                    flag = Pipe.pass_resume_text(value)
                if flag:
                    st.sidebar.success("document successfully uploaddeddd to vector store")
                else:
                    st.sidebar.warning("error in uploadinf file to vector store")
                
            elif file_type == "zip":
                st.sidebar.write("ZIP file uploaded")
                text = extract_text(uploaded_file,file_type)
                for key,value in text.items():
                    flag = Pipe.pass_resume_text(value)
                if flag:
                    st.sidebar.success("document successfully uploaddeddd to vector store")
                else:
                    st.sidebar.warning("error in uploadinf file to vector store")
            else:
                st.sidebar.write("Unsupported file type")

if 'messages' not in st.session_state:
    st.session_state.messages = []
    st.session_state.memory = InMemorySaver()
    st.session_state.thread = {"configurable": {"thread_id": "1"}}
    st.session_state.abot = Agent(st.session_state.memory)
    print("\033[92mAgent initialized\033[0m")

abot = st.session_state.abot

for messages in st.session_state.messages:
    with st.chat_message(messages["role"]):
        st.markdown(messages["content"])

if prompt:= st.chat_input("Enter your response"):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.spinner("Generating response", show_time=True):
        result = abot.graph.invoke({"user_input":[prompt]}, st.session_state.thread)
    
    with st.chat_message("assistant"):
        response = st.write_stream(response_generator(result['msg'][-1].content))

    st.session_state.messages.append({"role": "assistant", "content": response})